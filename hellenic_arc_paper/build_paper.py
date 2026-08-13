# -*- coding: utf-8 -*-
"""בונה את קובץ ה-Word של עבודת הסיכום על הקשת ההלנית.

דרישות העיצוב (מתוך הנחיות התרגיל):
    גוף המסמך  - Arial 12, מרווח שורה וחצי, יישור דו-צידי, כיווניות מימין לשמאל
    רשימת מקורות - Arial 11, מרווח שורה בודד, סדר אלפביתי
"""

import os

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCX_NAME = "עבודת_סיכום_הקשת_ההלנית.docx"

FONT = "Arial"
BODY_SIZE = Pt(12)
CAPTION_SIZE = Pt(11)
REF_SIZE = Pt(11)


# ----------------------------------------------------------------------------
# עזרי כיווניות (RTL) - python-docx אינו חושף אותם ישירות.
# סדר האלמנטים בתוך pPr/rPr/sectPr נאכף בסכמה של OOXML, ולכן אי אפשר פשוט
# לצרף אלמנט בסוף: קובץ עם סדר שגוי נדחה בפתיחה.
# ----------------------------------------------------------------------------
PPR_SEQ = (
    "w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore", "w:framePr",
    "w:widowControl", "w:numPr", "w:suppressLineNumbers", "w:pBdr", "w:shd",
    "w:tabs", "w:suppressAutoHyphens", "w:kinsoku", "w:wordWrap",
    "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE", "w:autoSpaceDN",
    "w:bidi", "w:adjustRightInd", "w:snapToGrid", "w:spacing", "w:ind",
    "w:contextualSpacing", "w:mirrorIndents", "w:suppressOverlap", "w:jc",
    "w:textDirection", "w:textAlignment", "w:textboxTightWrap", "w:outlineLvl",
    "w:divId", "w:cnfStyle", "w:rPr", "w:sectPr", "w:pPrChange",
)

RPR_SEQ = (
    "w:rStyle", "w:rFonts", "w:b", "w:bCs", "w:i", "w:iCs", "w:caps",
    "w:smallCaps", "w:strike", "w:dstrike", "w:outline", "w:shadow",
    "w:emboss", "w:imprint", "w:noProof", "w:snapToGrid", "w:vanish",
    "w:webHidden", "w:color", "w:spacing", "w:w", "w:kern", "w:position",
    "w:sz", "w:szCs", "w:highlight", "w:u", "w:effect", "w:bdr", "w:shd",
    "w:fitText", "w:vertAlign", "w:rtl", "w:cs", "w:em", "w:lang",
    "w:eastAsianLayout", "w:specVanish", "w:oMath",
)

SECTPR_SEQ = (
    "w:footnotePr", "w:endnotePr", "w:type", "w:pgSz", "w:pgMar", "w:paperSrc",
    "w:pgBorders", "w:lnNumType", "w:pgNumType", "w:cols", "w:formProt",
    "w:vAlign", "w:noEndnote", "w:titlePg", "w:textDirection", "w:bidi",
    "w:rtlGutter", "w:docGrid", "w:printerSettings", "w:sectPrChange",
)


def ensure_child(parent, tag, seq):
    """מוסיף אלמנט ריק במיקום התקין לפי סדר הסכמה (או מחזיר קיים)."""
    existing = parent.find(qn(tag))
    if existing is not None:
        return existing
    element = parent.makeelement(qn(tag), {})
    parent.insert_element_before(element, *seq[seq.index(tag) + 1:])
    return element


def set_paragraph_rtl(paragraph):
    """מסמן את הפסקה ככיוונית מימין לשמאל."""
    ensure_child(paragraph._p.get_or_add_pPr(), "w:bidi", PPR_SEQ)


def set_run_font(run, name=FONT, size=BODY_SIZE, rtl=True, bold=False):
    """קובע גופן לרוץ, כולל גופן complex-script - בלעדיו Word מתעלם מהגופן בעברית."""
    run.font.name = name
    run.font.size = size
    run.font.bold = bold
    r_pr = run._r.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        r_fonts.set(qn(attr), name)
    # גודל complex-script נפרד מגודל הגופן הלטיני
    ensure_child(r_pr, "w:szCs", RPR_SEQ).set(qn("w:val"), str(int(size.pt * 2)))
    if rtl:
        ensure_child(r_pr, "w:rtl", RPR_SEQ)
    if bold:
        ensure_child(r_pr, "w:bCs", RPR_SEQ)


def set_section_rtl(section):
    """מסמן את המקטע כמקטע מימין לשמאל."""
    ensure_child(section._sectPr, "w:bidi", SECTPR_SEQ)


def add_paragraph(doc, text, *, size=BODY_SIZE, bold=False, rtl=True,
                  align=WD_ALIGN_PARAGRAPH.JUSTIFY, line_spacing=1.5,
                  space_after=Pt(8), space_before=Pt(0)):
    p = doc.add_paragraph()
    p.alignment = align
    fmt = p.paragraph_format
    fmt.line_spacing = line_spacing
    fmt.space_after = space_after
    fmt.space_before = space_before
    if rtl:
        set_paragraph_rtl(p)
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, rtl=rtl, bold=bold)
    return p


def add_heading(doc, text):
    return add_paragraph(doc, text, bold=True,
                         align=WD_ALIGN_PARAGRAPH.RIGHT, line_spacing=1.5,
                         space_before=Pt(14), space_after=Pt(6))


def add_figure(doc, image_name, caption):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    set_paragraph_rtl(p)
    p.add_run().add_picture(os.path.join(OUT_DIR, image_name), width=Cm(15.5))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cap.paragraph_format.line_spacing = 1.0
    cap.paragraph_format.space_after = Pt(12)
    set_paragraph_rtl(cap)
    head, body = caption.split("|", 1)
    r1 = cap.add_run(head)
    set_run_font(r1, size=CAPTION_SIZE, bold=True)
    r2 = cap.add_run(body)
    set_run_font(r2, size=CAPTION_SIZE)
    return cap


def add_reference(doc, text):
    """פריט ברשימת המקורות - טקסט לטיני, ולכן פסקה משמאל לימין."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    fmt = p.paragraph_format
    fmt.line_spacing = 1.0
    fmt.space_after = Pt(6)
    fmt.left_indent = Cm(1.0)
    fmt.first_line_indent = Cm(-1.0)
    run = p.add_run(text)
    set_run_font(run, size=REF_SIZE, rtl=False)
    return p


# ----------------------------------------------------------------------------
# תוכן העבודה
# ----------------------------------------------------------------------------
TITLE = "מלח המסיני כמפתח למורפולוגיה של מנסרת ההצטברות בקשת ההלנית"
AUTHOR = "[שם מלא]"
COURSE = "עבודת סיכום המוגשת לקורס [שם ומספר הקורס]"

SECTIONS = [
    ("מבוא: מנסרת הצטברות ומפלס הניתוק שבבסיסה", [
        "מנסרת הצטברות היא גוף הסדימנטים שנגרד מעל הלוח השוקע ונערם על שוליו של "
        "הלוח העליון באזורי שקיעה. הגירוד אינו מתרחש בגובה שרירותי אלא לאורך מפלס "
        "ניתוק (décollement) – משטח חלש שלאורכו מתנתקים הסדימנטים העליונים מאלה "
        "שממשיכים לשקוע יחד עם הלוח. הצורה החיצונית של המנסרה נשלטת ישירות בחוזק של "
        "אותו מפלס: ככל שהוא חלש יותר, כך המנסרה שנבנית מעליו רחבה ושטוחה יותר. "
        "מכאן שהשאלה \"מדוע מנסרה מסוימת נראית כפי שהיא נראית\" היא בראש ובראשונה "
        "שאלה על החומר שבבסיסה.",

        "הקשת ההלנית שבמזרח הים התיכון מציעה את אחד המקרים הקיצוניים ביותר לבחינת "
        "הקשר הזה. מנסרת ההצטברות שלה, רכס הים התיכון, נבנתה מעל רצף של אידויים "
        "שהושקעו במשבר המליחות המסיניאני, כלומר מעל שכבת מלח – אחד החומרים החלשים "
        "ביותר שאפשר למצוא בבסיס טריז טקטוני. התוצאה היא מנסרה שאינה דומה לטריזים "
        "המוכרים משולי האוקיינוס השקט, לא בשיפועה, לא ברוחבה ולא בתופעות שעל פניה. "
        "עבודה זו בוחנת כיצד נוכחות המלח מסבירה את מורפולוגיית קרקעית הים באזור.",

        "כדי לענות על השאלה נדרש לשלב שלושה מישורי תיאור. תחילה יתואר האזור עצמו "
        "והמאפיינים הבתימטריים הבולטים בו, לאחר מכן תוצג הסביבה הגיאולוגית-טקטונית "
        "שמפעילה את המערכת, ולבסוף יידונו התהליכים – מכניים, הידרולוגיים "
        "וסדימנטריים – שמתרגמים את הטקטוניקה לצורות קרקעית. שלושת המישורים הללו "
        "מתלכדים סביב אותו גורם מארגן, והוא שכבת המלח המסיניאני.",
    ]),

    ("תיאור האזור: קשת, רכס ומערכת שקעים", [
        "הקשת ההלנית משתרעת מהים היוני שממערב לפלופונס, דרך דרום קרתה, ועד לאזור "
        "רודוס ודרום-מערב אנטוליה במזרח. המרכיב הבתימטרי הגדול ביותר בה הוא רכס הים "
        "התיכון – נפיחות קשתית, מוארכת ורחבה, החוצה את אגן מזרח הים התיכון לאורך "
        "כאלפיים קילומטרים מהאזור היוני ועד הלבנט (Huguen et al., 2006). הרכס אינו "
        "רכס וולקני אלא גוף תצורתי-טקטוני: הוא כולו ערימה של סדימנטים מקומטים "
        "ומועתקים (Kopf et al., 2003). כבר בהתבוננות ראשונה במפת עומקים ניכר "
        "שהמאפיין הבולט של האזור אינו שקע צר וחד, אלא נפיחות רחבה מאוד. איור 1 מציג "
        "את המערך הכללי של הקשת.",

        "צפונית לרכס נמצאת מערכת השקעים ההלניים, ובה השקעים מטפאן ופטולמאוס במערב "
        "והשקעים פליני וסטראבו במזרח. שקעים אלה הם החלקים העמוקים ביותר של הים "
        "התיכון כולו, ועומקם עולה על 5,000 מטרים (Mouslopoulou et al., 2025). "
        "הנקודה העמוקה ביותר, Calypso Deep, נמצאת בקצה המערבי של המערכת ועומקה כ-"
        "5,267 מטרים (Kontoyiannis et al., 2016). ריכוז העומקים המרביים דווקא כאן, "
        "ולא בקצה הדרומי של המערכת, הוא הרמז המורפולוגי הראשון לכך שהמערכת אינה "
        "בנויה כשקע אוקייני קלאסי.",

        "ההבחנה הזו חשובה דיה כדי להתעכב עליה. בשקע אוקייני טיפוסי, העומק המרבי "
        "נמצא בחזית העיוות, במקום שבו הלוח השוקע מתכופף ונכנס מתחת ללוח העליון. "
        "בקשת ההלנית, לעומת זאת, חזית העיוות שוכנת בקצה הדרומי של רכס הים התיכון, "
        "רחוק מן השקעים, והיא קבורה כמעט לחלוטין תחת כיסוי סדימנטרי עבה. השקעים "
        "עצמם הם אגנים בתוך הלוח העליון, בין המנסרה לבין הרכס הקדם-קשתי של קרתה "
        "(Mouslopoulou et al., 2025). מבחינה מורפולוגית, אפוא, המערכת מורכבת משני "
        "אלמנטים נפרדים – מנסרה רחבה בדרום ואגנים עמוקים בצפון – ולא מתעלה אחת.",
    ]),

    ("הסביבה הגיאולוגית-טקטונית: שקיעה, נסיגת לוח וקרע", [
        "המנוע של המערכת הוא שקיעת ליתוספרת נוביה צפונה, אל מתחת ללוח האגאי. "
        "התפלגות מוקדי רעידות האדמה בתוך הלוח השוקע מתווה את מיקומו עד לעומק של כ-160 "
        "קילומטרים (Shaw & Jackson, 2010). ההתכנסות נבלעת בעיקר לאורך העתקים הפוכים "
        "בעלי שיפוע מתון במגע שבין הלוחות, ובחלקה גם לאורך העתקי משנה תלולים יותר "
        "בתוך הלוח העליון. כלומר, מדובר במערכת שקיעה פעילה במלוא מובן המילה, אף שהיא "
        "יושבת בתוך אגן ים תיכוני סגור יחסית.",

        "הקצבים באזור חושפים ממצא מפתיע. קצב ההתכנסות בין הלוח האגאי לאפריקה הוא כ-35 "
        "מילימטרים בשנה, בעוד שקצב ההתכנסות בין נוביה לאירואסיה עומד על כ-6 מילימטרים "
        "בשנה בלבד (Shaw & Jackson, 2010). ההפרש אינו נובע מתנועת הלוחות הראשיים אלא "
        "מהתנועה המהירה של הלוח האגאי עצמו דרומה, תנועה המלווה בנסיגה של הלוח השוקע "
        "ובהתפשטות של האזור האגאי. במילים אחרות, המנסרה ניזונה מקצב התכנסות הגדול "
        "פי כמה מזה שהיו הלוחות הראשיים מספקים לבדם, וזהו תנאי ראשון לגדילתה המהירה.",

        "מאפיין נוסף של המגע בין הלוחות הוא צימוד נמוך. מרבית ההתכנסות, כ-80 אחוזים "
        "ממנה, נבלעת באופן א-סיסמי ולא בהחלקה סיסמית (Shaw & Jackson, 2010). ממצא זה "
        "מתיישב עם מגע לוחות חלש, שאינו מסוגל לאגור מאמצים גדולים לאורך זמן. הוא "
        "מצטרף לתמונה שתתחדד בהמשך: המערכת כולה, ממגע הלוחות ועד בסיס המנסרה, חלשה "
        "מכפי שמצופה ממערכת שקיעה בקצבים כאלה.",

        "המזרח של הקשת מתנהג אחרת מן המערב. באזור פליני-סטראבו ההתכנסות אלכסונית, "
        "ומחקרים מראים שהאזור כולו מתפקד כאזור גזירה שמאלי רחב שנוצר בעקבות קרע "
        "בליתוספרה השוקעת (Özbakır et al., 2013). הקרע עצמו מיוחס לנסיגה המתמשכת של "
        "הלוח, שיוצרת גזירה בין הליתוספרה האגאית לזו האפריקאית. לפיכך אין לצפות "
        "למורפולוגיה אחידה לאורך הקשת: אותה מערכת עוברת ממשטר של הצטברות חזיתית "
        "במערב למשטר של תנועת הזזה במזרח.",
    ]),

    ("מורפולוגיית המנסרה: רחבה, שטוחה, וגדלה במהירות", [
        "המאפיין המכני הבולט ביותר של רכס הים התיכון הוא זווית הטריז הנמוכה שלו. "
        "שיפוע פני האגף הדרומי של הרכס הוא בסדר גודל של מעלה אחת בלבד, וגם שיפוע "
        "הבסיס שמתחתיו דומה לו (Reston et al., 2002). על פי תורת הטריז הקריטי, "
        "שילוב כזה של שיפועים אפשרי רק כאשר מאמץ הגזירה בבסיס הטריז נמוך במיוחד. "
        "הגיאומטריה של הרכס, אם כן, היא בעצמה מדידה עקיפה של חולשת מפלס הניתוק. "
        "איור 2 מציג את היחסים המבניים הללו בחתך סכמטי.",

        "מהו אותו מפלס חלש? עבודות על חזית העיוות המערבית מראות שמפלס הניתוק ממוקם "
        "בתוך רצף האידויים המסיניאניים, ושנוכחותם היא שמכתיבה הן את השיפוע הנמוך והן "
        "את התקדמותו המהירה של ההעתק החזיתי קדימה (Reston et al., 2002). המלח חלש "
        "מסלעי סדימנט רגילים, והוא גם אינו מסוגל לשאת מאמצי גזירה גדולים לאורך זמן. "
        "נוסף על כך, אטימותו היחסית תורמת ללחצי נקבובים גבוהים בחומר שמתחתיו ומעליו, "
        "והדבר מחליש את המערכת עוד יותר. חולשת הבסיס אינה, אפוא, תכונה מקרית של "
        "האזור אלא תוצאה ישירה של אירוע גיאולוגי מוגדר בזמן.",

        "התוצאה הגיאומטרית של בסיס חלש היא מנסרה החורגת בממדיה מן המקובל. במקום טריז "
        "צר ותלול הנשען על שוליו של הלוח העליון, נוצר כאן גוף רחב מאוד המשתרע על פני "
        "מאות קילומטרים בכיוון ניצב לקשת. חזית העיוות שלו מתקדמת דרומה במהירות, עד "
        "כי רכס הים התיכון נחשב למנסרת ההצטברות הגדלה במהירות הרבה ביותר על פני כדור "
        "הארץ (Kopf et al., 2003). רוחב חריג וקצב גדילה חריג הם שני ביטויים של אותו "
        "גורם עצמו.",

        "בקצה הצפוני של המנסרה, במגע עם ה-backstop, המבנה משתנה. שם נבלמת המנסרה אל "
        "מול הליתוספרה הקשיחה של הלוח העליון, והעיוות עובר להעתקים הפוכים פעילים "
        "המשתרעים מן ה-backstop ועד לשקעים ההלניים (Mouslopoulou et al., 2025). "
        "כלומר, גם החלק ה\"פנימי\" של המערכת, זה שכבר אינו חלק מהטריז עצמו, נותר "
        "פעיל טקטונית. המעבר בין המנסרה לבין הלוח העליון הוא מעבר מבני חד, והוא בא "
        "לידי ביטוי בקרקעית הים כמדרון תלול היורד אל השקעים.",
    ]),

    ("תהליכים: הצטברות, לחץ-יתר, ובוצנות בוץ", [
        "לא כל הסדימנט הנכנס למערכת נערם במנסרה. מאזן מסה שנערך לאורך הקשת מראה חלוקה "
        "משתנה בין סדימנט המצטבר בטריז לבין סדימנט המוסע מתחתיו אל תוך אזור השקיעה, "
        "ומזהה שלושה משטרים שונים לאורך הקשת: קפיצה של מפלס הניתוק במערב, העתקה "
        "הפוכה אינטנסיבית במקטע המרכזי שבין לוב לקרתה, וטקטוניקה של תנועת הזזה במזרח "
        "(Kopf et al., 2003). כל משטר כזה מייצר חתימה מורפולוגית משלו בקרקעית. "
        "ההבדלים בין מקטעי הקשת אינם רעש, אלא ביטוי לשינוי בדרך שבה החומר עובר בין "
        "הלוחות.",

        "החומר שאינו מצטבר ואינו מוסע עד הסוף מוצא לעצמו מוצא שלישי – כלפי מעלה. "
        "בוץ רווי נוזלים ובעל לחץ-יתר, שמקורו באזור מפלס הניתוק, פורץ דרך הסדימנטים "
        "המעוותים של הרכס ומגיע עד לקרקעית הים, ושם הוא בונה הרי בוץ ודיאפירים. "
        "בשדה אולימפי, מדרום לקרתה, זוהו עשרות מבנים כאלה בעומקי מים של כ-1,700 עד "
        "2,000 מטרים (Camerlenghi et al., 1995). כלומר, החולשה ההידרולוגית של בסיס "
        "המנסרה אינה נשארת בעומק: היא מיתרגמת לצורות בולטות על פני הקרקעית.",

        "לתופעה יש גם מרכיב כימי מובהק. מי הנקבובים העולים עם הבוץ מלוחים בהרבה ממי "
        "הים, משום שבדרכם כלפי מעלה הם מגיבים עם האידויים המסיניאניים ומכילים את "
        "תוצרי המסתם. כאשר מים אלה מגיעים לקרקעית הם אינם מתערבבים עם מי הים אלא "
        "נאגרים בשקעים מקומיים ויוצרים בריכות מלוחה. בסקרי סונאר צדי הן מזוהות "
        "ככתמים נטולי הד, שאין להם הסבר טופוגרפי טוב יותר (Woodside & Volgin, 1996). "
        "המלח, שראשיתו כמפלס ניתוק מכני בעומק, מופיע כאן שוב כמעצב ישיר של סביבת "
        "הקרקעית.",

        "גם הטופוגרפיה העדינה של פני הרכס נושאת חתימה של אותם תהליכים. פני הרכס "
        "מאופיינים בטופוגרפיה גלית וסבוכה, מסוג \"אבני מרצפת\", ובשדות של שקעים "
        "מעוגלים המשתרעים לאורכו (Huguen et al., 2006). מבנים אלה נוצרים משילוב של "
        "קמטים וגלישות בכיסוי הסדימנטרי מעל שכבת המלח הזוחלת מתחתיו. כך מתקבלת "
        "מורפולוגיה שאינה נובעת מקימוט טקטוני בלבד אלא מזרימה איטית של החומר שבבסיס.",
    ]),

    ("תנועות סדימנט וזרמי קרקעית בשקעים", [
        "השקעים ההלניים מתפקדים כמלכודות סדימנט יעילות. מרבית מילויָם מורכב "
        "מטורבידיטים ומתצורות אחרות של זרימות כובד, ומיקומם של מוקדי השיקוע נשלט "
        "בהעתקים הפעילים באגן (Mouslopoulou et al., 2025). בשקע פליני, למשל, מסודרים "
        "מוקדי השיקוע בסדרה מדורגת שגיאומטרייתה משקפת את תנועת ההזזה השמאלית לאורך "
        "ציר השקע. הסדימנטציה, אפוא, אינה עצמאית: היא רושמת את המבנה הטקטוני שמתחתיה.",

        "העמקים הללו אינם גופי מים מנותקים. מדידות זרמים ומסות מים ב-Calypso Deep "
        "הראו עדשה של מים עמוקים ומלוחים בעומק של כ-3,000 מטרים, וכן חדירה אפיזודית "
        "של מים צפופים ומאווררים ממקור אדריאטי, שהעשירה בחמצן את השכבות התחתונות "
        "ביותר (Kontoyiannis et al., 2016). באותה תקופה השתנתה גם תבנית הזרימה בשקע, "
        "מסחרור אנטי-ציקלוני לסחרור ציקלוני מתחת לעומק של כ-2,500 מטרים. תחתית "
        "השקעים היא, אם כן, סביבה דינמית שמתחדשת מדי פעם, ולא אגן קפוא.",

        "לכיסוי הסדימנטרי יש גם השפעה מורפולוגית ישירה על החזית הדרומית של המערכת. "
        "חתך הסדימנט העבה שמצטבר על הלוח הנכנס ממסך את ביטויה הבתימטרי של חזית "
        "העיוות, עד כי היא כמעט אינה נראית כמדרגה בקרקעית. זו הסיבה שהמערכת אינה "
        "מציגה שקע צר וברור בקצה הדרומי שלה, כפי שמצופה משוליים מתכנסים. שיעור "
        "האספקה הסדימנטרית והמורפולוגיה הנצפית קשורים כאן זה בזה ישירות.",
    ]),

    ("סיכום ומסקנות", [
        "הקשת ההלנית פועלת בקצב התכנסות של כ-35 מילימטרים בשנה בין הלוח האגאי "
        "לאפריקה, קצב הגבוה בהרבה מן ההתכנסות בין הלוחות הראשיים. ההפרש נובע מנסיגת "
        "הלוח השוקע ומהתפשטות האזור האגאי, ולא מתנועה מהירה של אפריקה. קצב זה הוא "
        "התנאי הראשון להיווצרותה של מנסרה גדולה במיוחד.",

        "המורפולוגיה של האזור אינה זו של שקע אוקייני יחיד. חזית העיוות שוכנת בדרום "
        "וקבורה תחת סדימנט, ואילו העומקים המרביים – ובהם Calypso Deep, כ-5,267 "
        "מטרים – נמצאים באגנים שבתוך הלוח העליון, בין המנסרה לקרתה. מדובר, אפוא, "
        "בשתי יחידות מורפולוגיות שונות שיש להבחין ביניהן.",

        "מפלס הניתוק שבתוך האידויים המסיניאניים הוא ההסבר המרכזי לצורת המנסרה. חולשתו "
        "מאפשרת שיפועים בסדר גודל של מעלה אחת בלבד, רוחב חריג, וקצב התקדמות חזיתית "
        "שהופך את רכס הים התיכון למנסרת ההצטברות הגדלה במהירות הרבה ביותר בעולם. "
        "ללא שכבת המלח לא ניתן להסביר את הגיאומטריה הנצפית.",

        "אותו מפלס חלש ורווי נוזלים אחראי גם לתופעות הקרקעית הייחודיות. לחץ-יתר בבסיס "
        "מזין בוצנות בוץ והרי בוץ בשדה אולימפי, מי נקבובים מלוחים יוצרים בריכות מלוחה "
        "על פני הקרקעית, וזחילת המלח תורמת לטופוגרפיה הגלית של פני הרכס. הצורות "
        "הקטנות והצורה הגדולה נובעות מאותו גורם.",

        "לבסוף, המערכת אינה אחידה לאורכה. במערב שולטת הצטברות חזיתית, במקטע המרכזי "
        "העתקה הפוכה, ובמזרח – באזור פליני-סטראבו – עוברת המערכת למשטר של תנועת הזזה "
        "בעקבות קרע בלוח השוקע. לפיכך יש לקרוא את מורפולוגיית הקרקעית של הקשת ההלנית "
        "כרצף של מקטעים, שכל אחד מהם מבטא את האופן שבו החומר עובר בין הלוחות באותו "
        "מקטע.",
    ]),
]

FIGURES = {
    # מוצב אחרי הפסקה הראשונה של הפרק המצוין
    "תיאור האזור: קשת, רכס ומערכת שקעים": (
        3, "figure1_map.png",
        "איור 1.|"
        " מפה סכמטית של הקשת ההלנית. חשוב לראות בה שלושה אלמנטים ואת היחס ביניהם: "
        "חזית העיוות (קו אדום) בדרום, רצועת רכס הים התיכון – מנסרת ההצטברות – "
        "שמצפון לה, ומערכת השקעים ההלניים (קו כחול) הנמצאת צפונית לרכס ודרומית "
        "לקרתה. שימו לב שהעומק המרבי (Calypso Deep) נמצא בשקעים ולא בחזית העיוות, "
        "ושהרי הבוץ מרוכזים על האגף הצפוני של הרכס. החץ מציין את כיוון תנועת נוביה "
        "ביחס ללוח האגאי. שורטט על בסיס הנתונים המתוארים אצל "
        "Huguen et al. (2006), Shaw & Jackson (2010) ו-Mouslopoulou et al. (2025). "
        "המפה סכמטית ואינה מדויקת בקנה מידה."
    ),
    "מורפולוגיית המנסרה: רחבה, שטוחה, וגדלה במהירות": (
        1, "figure2_cross_section.png",
        "איור 2.|"
        " חתך רוחב סכמטי דרך המערכת, מדרום (משמאל) לצפון (מימין). חשוב לראות את "
        "מפלס הניתוק (קו אדום) הממוקם בתוך שכבת המלח המסיניאני (צהוב), ואת השיפוע "
        "המתון מאוד של פני המנסרה שנובע ממנו. מעל מפלס הניתוק מצטברים הסדימנטים "
        "בהעתקים הפוכים; מתחתיו ממשיך הרצף הטרום-מסיניני ולוח נוביה לשקוע צפונה. "
        "החץ הכחול מסמן את מסלול הנוזלים בעלי לחץ-היתר שמזינים את הרי הבוץ שעל "
        "הקרקעית. מצפון למנסרה נמצאים ה-backstop והשקע ההלני. שורטט על בסיס "
        "התיאורים אצל Reston et al. (2002), Kopf et al. (2003) ו-"
        "Mouslopoulou et al. (2025). החתך סכמטי, והממדים האנכיים מוגזמים."
    ),
}

REFERENCES = [
    "Camerlenghi, A., Cita, M. B., Della Vedova, B., Fusi, N., Mirabile, L., & "
    "Pellis, G. (1995). Geophysical evidence of mud diapirism on the "
    "Mediterranean Ridge accretionary complex. Marine Geophysical Researches, "
    "17, 115–141.",

    "Huguen, C., Chamot-Rooke, N., Loubrieu, B., & Mascle, J. (2006). "
    "Morphology of a pre-collisional, salt-bearing, accretionary complex: the "
    "Mediterranean Ridge (Eastern Mediterranean). Marine Geophysical "
    "Researches, 27(1), 61–75.",

    "Kontoyiannis, H., Lykousis, V., Papadopoulos, V., Stavrakakis, S., "
    "Anassontzis, E. G., Belias, A., Koutsoukos, S., & Resvanis, L. K. (2016). "
    "Hydrography, "
    "circulation, and mixing at the Calypso Deep (the deepest Mediterranean "
    "trough) during 2006–09. Journal of Physical Oceanography, 46(4), "
    "1255–1276.",

    "Kopf, A., Mascle, J., & Klaeschen, D. (2003). The Mediterranean Ridge: a "
    "mass balance across the fastest growing accretionary complex on Earth. "
    "Journal of Geophysical Research: Solid Earth, 108(B8), 2372.",

    "Mouslopoulou, V., Begg, J. G., Polonia, A., Nicol, A., Reston, T. J., "
    "Cesca, S., Giba, M., & Gasperini, L. (2025). Hellenic subduction system "
    "and upper-plate "
    "structures revealed by deep high-resolution seismic-reflection profiles "
    "and seafloor bathymetry. Tectonics, 44, e2025TC008943.",

    "Özbakır, A. D., Şengör, A. M. C., Wortel, M. J. R., & Govers, R. (2013). "
    "The Pliny–Strabo trench region: a large shear zone resulting from slab "
    "tearing. Earth and Planetary Science Letters, 375, 188–195.",

    "Reston, T. J., von Huene, R., Dickmann, T., Klaeschen, D., & Kopp, H. "
    "(2002). Frontal accretion along the western Mediterranean Ridge: the "
    "effect of Messinian evaporites on wedge mechanics and structural style. "
    "Marine Geology, 186, 59–82.",

    "Shaw, B., & Jackson, J. (2010). Earthquake mechanisms and active tectonics "
    "of the Hellenic subduction zone. Geophysical Journal International, "
    "181(2), 966–984.",

    "Woodside, J. M., & Volgin, A. V. (1996). Brine pools associated with "
    "Mediterranean Ridge mud diapirs: an interpretation of echo-free patches "
    "in deep tow sidescan sonar data. Marine Geology, 132, 55–61.",
]


def build():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)
    set_section_rtl(section)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal.font.size = BODY_SIZE
    normal.element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.element.rPr.rFonts.set(qn("w:eastAsia"), FONT)

    # --- ראש המסמך ---
    add_paragraph(doc, TITLE, bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT,
                  space_after=Pt(2))
    add_paragraph(doc, AUTHOR, align=WD_ALIGN_PARAGRAPH.RIGHT,
                  space_after=Pt(2))
    add_paragraph(doc, COURSE, align=WD_ALIGN_PARAGRAPH.RIGHT,
                  space_after=Pt(2))
    add_paragraph(doc, "", space_after=Pt(0))  # שורת רווח

    # --- גוף המסמך ---
    for title, paragraphs in SECTIONS:
        add_heading(doc, title)
        fig = FIGURES.get(title)
        for i, text in enumerate(paragraphs, start=1):
            add_paragraph(doc, text)
            if fig and fig[0] == i:
                add_figure(doc, fig[1], fig[2])

    # --- רשימת מקורות ---
    add_heading(doc, "רשימת מקורות")
    for ref in REFERENCES:
        add_reference(doc, ref)

    path = os.path.join(OUT_DIR, DOCX_NAME)
    doc.save(path)
    return path


if __name__ == "__main__":
    print(build())

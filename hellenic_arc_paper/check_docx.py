# -*- coding: utf-8 -*-
"""בדיקת תקינות של קובץ העבודה מול דרישות העיצוב של התרגיל."""

import os
import unicodedata
import sys

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                    "עבודת_סיכום_הקשת_ההלנית.docx")


def has(el, tag):
    return el is not None and el.find(qn(tag)) is not None


def main():
    doc = Document(PATH)
    problems = []
    body_words = 0
    body_paras = 0
    ref_paras = 0
    short_paras = []

    for p in doc.paragraphs:
        text = p.text.strip()
        if not text:
            continue
        p_pr = p._p.pPr
        is_ref = p.paragraph_format.first_line_indent is not None and \
            p.paragraph_format.first_line_indent < 0

        for run in p.runs:
            r_pr = run._r.rPr
            fonts = r_pr.find(qn("w:rFonts")) if r_pr is not None else None
            if fonts is None or fonts.get(qn("w:cs")) != "Arial":
                problems.append(f"רוץ ללא Arial ב-complex-script: {text[:40]}")
            if run.font.name != "Arial":
                problems.append(f"רוץ ללא Arial: {text[:40]}")
            if not is_ref and not has(r_pr, "w:rtl"):
                problems.append(f"רוץ עברי ללא w:rtl: {text[:40]}")

        if is_ref:
            ref_paras += 1
            if p.runs and p.runs[0].font.size != Pt(11):
                problems.append(f"מקור שאינו בגודל 11: {text[:40]}")
            if p.paragraph_format.line_spacing != 1.0:
                problems.append(f"מקור שאינו במרווח שורה בודד: {text[:40]}")
            continue

        if not has(p_pr, "w:bidi"):
            problems.append(f"פסקה ללא bidi: {text[:40]}")

        size = p.runs[0].font.size if p.runs else None
        if size == Pt(12) and p.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY:
            body_paras += 1
            body_words += len(text.split())
            if text.count(".") < 3:
                short_paras.append(text[:60])
            if p.paragraph_format.line_spacing != 1.5:
                problems.append(f"פסקת גוף שאינה במרווח 1.5: {text[:40]}")

    images = len(doc.inline_shapes)

    print(f"פסקאות גוף (Arial 12, justified): {body_paras}")
    print(f"מילים בגוף העבודה: {body_words}")
    print(f"פריטים ברשימת המקורות: {ref_paras}")
    print(f"איורים מוטמעים: {images}")
    print(f"הערכת עמודים (≈330 מילים/עמוד + איורים): "
          f"{body_words / 330 + images * 0.55:.1f}")

    refs_sorted = []
    for p in doc.paragraphs:
        if p.paragraph_format.first_line_indent is not None and \
                p.paragraph_format.first_line_indent < 0:
            refs_sorted.append(p.text.strip())
    def sort_key(s):
        """סדר אלפביתי לטיני - מתעלם מסימני ניקוד (Ö נחשב O)."""
        stripped = unicodedata.normalize("NFKD", s)
        return "".join(c for c in stripped
                       if not unicodedata.combining(c)).lower()

    if refs_sorted != sorted(refs_sorted, key=sort_key):
        problems.append("רשימת המקורות אינה בסדר אלפביתי")

    if short_paras:
        print("\nפסקאות עם פחות משלושה משפטים:")
        for s in short_paras:
            print("  -", s)

    if problems:
        print("\nבעיות שנמצאו:")
        for pr in sorted(set(problems)):
            print("  -", pr)
        return 1
    print("\nכל בדיקות העיצוב עברו.")
    return 0


if __name__ == "__main__":
    sys.exit(main())

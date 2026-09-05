# -*- coding: utf-8 -*-
"""בונה מסמך אימות נפרד לכל הערות השוליים שבעבודה.

לכל הערה: מספר, מקור, לוקטור, הטענה כפי שנוסחה בעבודה, והציטוט המילולי
מן המקור. כל ציטוט מאומת קודם מול קובץ ה-PDF עצמו - מול העמוד המודפס
(Huguen, Güneş) או מול פסקת ה-AGU (Kopf, Royden) שהלוקטור טוען לו.
ציטוט שאינו נמצא שם מפיל את הבנייה.

הרצה:  python3 build_verification.py
"""

import os
import re
import sys
import unicodedata
import zipfile

import fitz
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from lxml import etree

from quotes import QUOTES, ATTRIBUTION

HERE = os.path.dirname(os.path.abspath(__file__))
UPLOADS = "/root/.claude/uploads/f8b1b0e4-1ae6-519f-96d8-6c82ed280b6b"
DOC = os.path.join(HERE, "עבודת_סיכום_הקשת_ההלנית.docx")
OUT = os.path.join(HERE, "טבלת_אימות_ציטוטים.docx")

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

PDFS = {
    "huguen": "426c45fb-HuguenChamotRookeMas2006.pdf",
    "kopf": None,       # מזוהה לפי תבנית שם
    "gunes": None,
    "royden": None,
}
PATTERNS = {
    "huguen": "Huguen",
    "kopf": "Kopf",
    "gunes": "Gu_nes",
    "royden": "Royden",
}
# היסט בין מספר העמוד המודפס למספר העמוד בקובץ ה-PDF
PAGE_OFFSET = {"huguen": 61, "gunes": 297}

SHORT = {
    "huguen": "Huguen et al. (2006)",
    "kopf": "Kopf et al. (2003)",
    "gunes": "Güneş et al. (2018)",
    "royden": "Royden & Papanikolaou (2011)",
}

FULL_REFS = [
    "Güneş, P., Aksu, A. E., & Hall, J. (2018). Internal seismic stratigraphy "
    "of the Messinian evaporites across the northern sector of the eastern "
    "Mediterranean Sea. Marine and Petroleum Geology, 91, 297–320.",
    "Huguen, C., Chamot-Rooke, N., Loubrieu, B., & Mascle, J. (2006). "
    "Morphology of a pre-collisional, salt-bearing, accretionary complex: The "
    "Mediterranean Ridge (Eastern Mediterranean). Marine Geophysical "
    "Researches, 27(1), 61–75.",
    "Kopf, A., Mascle, J., & Klaeschen, D. (2003). The Mediterranean Ridge: A "
    "mass balance across the fastest growing accretionary complex on Earth. "
    "Journal of Geophysical Research: Solid Earth, 108(B8), 2372.",
    "Royden, L. H., & Papanikolaou, D. J. (2011). Slab segmentation and late "
    "Cenozoic disruption of the Hellenic arc. Geochemistry, Geophysics, "
    "Geosystems, 12(3), Q03010.",
]

# משפט הטענה של שתי הערות אינו ניתן לחילוץ אוטומטי: האחת היא ציטוט נרטיבי
# שנפרש על שני משפטים, והשנייה היא שורת המקור שבכיתוב איור 2.
CLAIM_OVERRIDE = {
    # הציטוט מגבה גם את משפט הפתיחה של הפסקה, ולא רק את המשפט שאליו הוצמד
    # סימן ההערה.
    36: "מאזן המסה של המנסרה נבחן באמצעות איזון נפחי בין כמות המשקעים הנצברת "
        "לזו הנבלעת בהפחתה. הניתוח התבסס על ארבעה חתכים סיסמיים לרוחב המנסרה, "
        "שעובדו לחתכי עומק אמיתיים.",
}


# --------------------------------------------------------------- אימות ----

def norm(s):
    """השוואה עמידה לשבירת שורות, ליגטורות, מקפים וסימני פיסוק טיפוגרפיים."""
    s = unicodedata.normalize("NFKD", s).lower()
    return re.sub(r"[^a-z0-9%]", "", s)


def find_pdf(key):
    pat = PATTERNS[key]
    for name in sorted(os.listdir(UPLOADS)):
        if pat in name and name.lower().endswith(".pdf"):
            return os.path.join(UPLOADS, name)
    raise SystemExit(f"לא נמצא PDF עבור {key}")


class Source:
    def __init__(self, key):
        self.key = key
        self.doc = fitz.open(find_pdf(key))
        self.pages = [p.get_text() for p in self.doc]
        self.raw = "\n".join(self.pages)
        # מיקומי הפסקאות בפורמט AGU, לפי סדר רץ
        self.para = {}
        expected = 1
        for m in re.finditer(r"\[(\d{1,3})\]", self.raw):
            if int(m.group(1)) == expected:
                self.para[expected] = m.start()
                expected += 1
        self.last_para = expected - 1

    def page_text(self, printed):
        idx = printed - PAGE_OFFSET[self.key]
        if not 0 <= idx < len(self.pages):
            raise SystemExit(f"{self.key}: עמוד {printed} מחוץ לתחום")
        return self.pages[idx], idx + 1

    def para_text(self, n):
        if n not in self.para:
            raise SystemExit(f"{self.key}: פסקה {n} לא אותרה")
        start = self.para[n]
        end = self.para.get(n + 1, len(self.raw))
        return self.raw[start:end]

    def page_of(self, offset_text):
        """מספר העמוד בקובץ שבו מופיע קטע טקסט."""
        pos, run = 0, 0
        for i, page in enumerate(self.pages):
            run += len(page) + 1
            if offset_text < run:
                return i + 1
            pos = run
        return len(self.pages)


def verify(sources, nid, src, loc, quote):
    s = sources[src]
    nq = norm(quote)
    if src in PAGE_OFFSET:
        text, pdf_page = s.page_text(loc)
        where = f"עמ' {loc} (עמוד {pdf_page} בקובץ ה-PDF)"
        ok = nq in norm(text)
    elif loc == "fig1a":
        hit = s.raw.find(quote[:40])
        if hit < 0:
            hit = None
        ok = nq in norm(s.raw)
        pdf_page = s.page_of(hit) if hit is not None else "?"
        where = f"כיתוב איור 1a (עמוד {pdf_page} בקובץ ה-PDF)"
    else:
        text = s.para_text(loc)
        pdf_page = s.page_of(s.para[loc])
        where = f"פס' {loc} (עמוד {pdf_page} בקובץ ה-PDF)"
        ok = nq in norm(text)
    return ok, where


# ------------------------------------------------- קריאת ההערות מן העבודה ----

def read_paper():
    z = zipfile.ZipFile(DOC)
    notes = etree.fromstring(z.read("word/footnotes.xml"))
    doc = etree.fromstring(z.read("word/document.xml"))
    texts = {}
    for note in notes.findall(f"{W}footnote"):
        nid = int(note.get(f"{W}id"))
        if nid > 0:
            texts[nid] = "".join(
                t.text or "" for t in note.findall(f".//{W}t")).strip()
    claims = {}
    for para in doc.findall(f".//{W}p"):
        acc = ""
        for node in para.iter():
            if node.tag == f"{W}t":
                acc += node.text or ""
            elif node.tag == f"{W}footnoteReference":
                nid = int(node.get(f"{W}id"))
                parts = re.split(r"(?<=[.!?])\s+", acc.strip())
                claims[nid] = parts[-1].strip() if parts else ""
    return texts, claims


# ------------------------------------------------------------- בניית docx ----

PPR_SEQ = ("w:pStyle", "w:keepNext", "w:keepLines", "w:pageBreakBefore",
           "w:framePr", "w:widowControl", "w:numPr", "w:suppressLineNumbers",
           "w:pBdr", "w:shd", "w:tabs", "w:suppressAutoHyphens", "w:kinsoku",
           "w:wordWrap", "w:overflowPunct", "w:topLinePunct", "w:autoSpaceDE",
           "w:autoSpaceDN", "w:bidi", "w:adjustRightInd", "w:snapToGrid",
           "w:spacing", "w:ind", "w:contextualSpacing", "w:mirrorIndents",
           "w:suppressOverlap", "w:jc", "w:textDirection", "w:textAlignment",
           "w:textboxTightWrap", "w:outlineLvl", "w:divId", "w:cnfStyle",
           "w:rPr", "w:sectPr", "w:pPrChange")
RPR_SEQ = ("w:rStyle", "w:rFonts", "w:b", "w:bCs", "w:i", "w:iCs", "w:caps",
           "w:smallCaps", "w:strike", "w:dstrike", "w:outline", "w:shadow",
           "w:emboss", "w:imprint", "w:noProof", "w:snapToGrid", "w:vanish",
           "w:webHidden", "w:color", "w:spacing", "w:w", "w:kern",
           "w:position", "w:sz", "w:szCs", "w:highlight", "w:u", "w:effect",
           "w:bdr", "w:shd", "w:fitText", "w:vertAlign", "w:rtl", "w:cs",
           "w:em", "w:lang", "w:eastAsianLayout", "w:specVanish", "w:oMath")
SECTPR_SEQ = ("w:footnotePr", "w:endnotePr", "w:type", "w:pgSz", "w:pgMar",
              "w:paperSrc", "w:pgBorders", "w:lnNumType", "w:pgNumType",
              "w:cols", "w:formProt", "w:vAlign", "w:noEndnote",
              "w:titlePg", "w:textDirection", "w:bidi", "w:rtlGutter",
              "w:docGrid", "w:printerSettings", "w:sectPrChange")


def ensure(parent, tag, seq):
    el = parent.find(qn(tag))
    if el is None:
        el = parent.makeelement(qn(tag), {})
        idx = seq.index(tag)
        parent.insert_element_before(el, *seq[idx + 1:])
    return el


def fmt(par, *, rtl=True, bold=False, italic=False, size=11,
        left=0.0, space_after=2, keep=False):
    ppr = par._p.get_or_add_pPr()
    if rtl:
        ensure(ppr, "w:bidi", PPR_SEQ)
    if keep:
        ensure(ppr, "w:keepNext", PPR_SEQ)
    sp = ensure(ppr, "w:spacing", PPR_SEQ)
    sp.set(qn("w:line"), "240")
    sp.set(qn("w:lineRule"), "auto")
    sp.set(qn("w:after"), str(int(space_after * 20)))
    sp.set(qn("w:before"), "0")
    if left:
        ind = ensure(ppr, "w:ind", PPR_SEQ)
        val = str(int(left * 567))
        for name in (("w:right", "w:end") if rtl else ("w:left", "w:start")):
            ind.set(qn(name), val)
    jc = ensure(ppr, "w:jc", PPR_SEQ)
    jc.set(qn("w:val"), "right" if rtl else "left")
    for run in par.runs:
        rpr = run._r.get_or_add_rPr()
        fonts = ensure(rpr, "w:rFonts", RPR_SEQ)
        for a in ("w:ascii", "w:hAnsi", "w:cs"):
            fonts.set(qn(a), "Arial")
        for tag, on in (("w:b", bold), ("w:bCs", bold),
                        ("w:i", italic), ("w:iCs", italic)):
            if on:
                ensure(rpr, tag, RPR_SEQ)
        for tag in ("w:sz", "w:szCs"):
            ensure(rpr, tag, RPR_SEQ).set(qn("w:val"), str(int(size * 2)))
        ensure(rpr, "w:rtl", RPR_SEQ).set(qn("w:val"), "1" if rtl else "0")
        lang = ensure(rpr, "w:lang", RPR_SEQ)
        lang.set(qn("w:bidi"), "he-IL")
        lang.set(qn("w:val"), "en-US")
    return par


def main():
    sources = {k: Source(k) for k in ("huguen", "kopf", "gunes", "royden")}
    texts, claims = read_paper()

    checked, failures, located = 0, [], {}
    for nid in sorted(texts):
        if nid in ATTRIBUTION:
            continue
        if nid not in QUOTES:
            failures.append(f"הערה {nid}: אין ציטוט מוגדר")
            continue
        located[nid] = []
        for src, loc, quote in QUOTES[nid]:
            ok, where = verify(sources, nid, src, loc, quote)
            checked += 1
            if not ok:
                failures.append(f"הערה {nid}: הציטוט לא נמצא ב{where} "
                                f"אצל {SHORT[src]}")
            located[nid].append((src, where, quote))

    print(f"ציטוטים שנבדקו: {checked}")
    if failures:
        print("\nכשלי אימות:")
        for f in failures:
            print("  -", f)
        return 1
    print("כל הציטוטים אומתו מול קובצי ה-PDF.\n")

    # ---------------------------------------------------------- המסמך ----
    doc = Document()
    sec = doc.sections[0]
    sec.page_height, sec.page_width = Cm(29.7), Cm(21.0)
    for side in ("left_margin", "right_margin"):
        setattr(sec, side, Cm(2.0))
    sec.top_margin = sec.bottom_margin = Cm(2.0)
    spr = sec._sectPr
    ensure(spr, "w:bidi", SECTPR_SEQ).set(qn("w:val"), "1")

    st = doc.styles["Normal"]
    st.font.name, st.font.size = "Arial", Pt(11)

    fmt(doc.add_paragraph("מסמך אימות ציטוטים"), bold=True, size=14,
        space_after=4, keep=True)
    fmt(doc.add_paragraph(
        "מסמך עבודה פנימי, נלווה לעבודת הסיכום \"מערכות שקע אוקייני ומנסרות הצטברות: "
        "הקשת ההלנית\". המסמך אינו חלק מן ההגשה."), size=11, space_after=8)
    fmt(doc.add_paragraph(
        "לכל אחת מהערות השוליים שבעבודה מובאים כאן המקור, הלוקטור המדויק "
        "(עמוד מודפס אצל Huguen ו-Güneş; מספר פסקה בפורמט AGU אצל Kopf "
        "ו-Royden), הטענה כפי שנוסחה בעבודה, והציטוט המילולי שממנו היא נגזרת. "
        f"כל {checked} הציטוטים אותרו אוטומטית בקובצי ה-PDF של המאמרים, בעמוד "
        "או בפסקה שהלוקטור טוען לו, לפני שנכתבו לכאן."), space_after=10)

    fmt(doc.add_paragraph("המקורות"), bold=True, size=12,
        space_after=4, keep=True)
    for ref in FULL_REFS:
        fmt(doc.add_paragraph(ref), rtl=False, size=10, left=0.6,
            space_after=4)

    fmt(doc.add_paragraph(""), space_after=6)
    fmt(doc.add_paragraph("ההערות"), bold=True, size=12,
        space_after=6, keep=True)

    for nid in sorted(texts):
        if nid in ATTRIBUTION:
            fmt(doc.add_paragraph(f"הערה {nid} — ייחוס מקור"),
                bold=True, size=11, space_after=2, keep=True)
            fmt(doc.add_paragraph(f"תוכן ההערה: {texts[nid]}"),
                left=0.5, space_after=2)
            fmt(doc.add_paragraph(ATTRIBUTION[nid]), left=0.5, space_after=10)
            continue

        first_src = located[nid][0][0]
        fmt(doc.add_paragraph(f"הערה {nid} — {SHORT[first_src]}"),
            bold=True, size=11, space_after=2, keep=True)
        fmt(doc.add_paragraph(f"הערת השוליים בעבודה: {texts[nid]}"),
            left=0.5, space_after=2)
        claim = CLAIM_OVERRIDE.get(nid, claims.get(nid, ""))
        fmt(doc.add_paragraph(f"הטענה בעבודה: {claim}"),
            left=0.5, space_after=2)
        for src, where, quote in located[nid]:
            label = "ציטוט מדויק" if len(located[nid]) == 1 \
                else f"ציטוט מדויק ({SHORT[src]})"
            fmt(doc.add_paragraph(f"{label} — אומת ב{where}:"),
                left=0.5, space_after=1, keep=True)
            fmt(doc.add_paragraph(f"“{quote}”"),
                rtl=False, italic=True, size=10, left=1.2, space_after=2)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(10)

    doc.save(OUT)
    print(OUT)
    return 0


if __name__ == "__main__":
    sys.exit(main())
